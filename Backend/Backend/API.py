from flask import Flask, request
from docx import Document

app = Flask(__name__)

@app.route('/find-sections', methods=['POST'])
def find_sections():
    # Check for missing request data
    if 'file_path' not in request.json:
        return 'Missing file_path field', 400
    if 'search_terms' not in request.json:
        return 'Missing search_terms field', 400
    if 'sections' not in request.json:
        return 'Missing sections field', 400
    if 'lines' not in request.json:
        return 'Missing lines field', 400

    # Get the file path from the request
    file_path = request.json['file_path']

    # Read the file
    with open(file_path, 'r') as f:
        lines = f.readlines()

    # Get the search terms, sections, and lines from the request
    search_terms = request.json['search_terms']
    sections = request.json['sections']
    lines = request.json['lines']

    # Create a new document
    document = Document()

    # Iterate over the search terms and find the corresponding sections
    for term in search_terms:
        lineNo = 0
        termLineNo = []
        lineEmpty = 0
        termsNum = 0
        for line in Lines:
            if term in line:
                termLineNo.append(lineNo)
                termsNum += 1
                print(termsNum, term)
            lineNo += 1


        # Add the sections to the document
        for i in sections:
            section_lines = lines[i-1].split(":")
            start_line = term_line_no[i-1]
            line_empty = 0

            if section_lines[0] == 'WHOLE':
                while line_empty == 0:
                    if lines[start_line] != "\n":
                        section = document.add_paragraph(lines[start_line])
                        start_line += 1
                    else:
                        line_empty = 1

            elif section_lines[0] == 'FIRST':
                line_count = 0
                while line_count < int(section_lines[1]):
                    section = document.add_paragraph(lines[start_line])
                    start_line += 1
                    line_count += 1

            elif section_lines[0] == 'LAST':
                line_count = 0
                while line_count < int(section_lines[1]):
                    section = document.add_paragraph(lines[start_line+10])
                    start_line += 1
                    line_count += 1

            elif section_lines[0] == 'SPECIFIC':
                specific_lines = [int(l) for l in section_lines[1].split(",")]
                document.add_paragraph(lines[start_line])
                for l in specific_lines:
                    section = document.add_paragraph(lines[start_line + l + 1])

    try:
        # Save the document
        document.save("data_conversion.docx")
    except Exception as e:
        return f'Error saving document: {e}', 500

    return 'OK'

if __name__ == '__main__':
    app.run(debug=True)


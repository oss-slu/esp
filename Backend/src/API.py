from flask import Flask, request, jsonify
from docx import Document
from pathlib import Path
from flask_cors import CORS
import json

app = Flask(__name__)
CORS(app)

@app.route('/find-sections', methods=['POST'])
def find_sections():
    

    # Check for missing request data
    if 'file' not in request.files:
        return 'No file part', 400
    

    # Get the file path from the request
    file = request.files['file']


    # if user does not select file, browser submits an empty file without a filename.
    if file.filename == '':
        return 'No selected file', 400
    

    
    #Ensures the additional data sent is properly received. 
    data_str = request.form.get('data', None)
    if data_str is None:
        return 'Missing additional data', 400    
    try:
        data_json = json.loads(data_str)
    except json.JSONDecodeError:
        return 'Invalid JSON data', 400
        
    # Get the search terms, sections, and lines from the request 
    if len(data_json) != 5:
        return 'Missing one or more required fields', 400
    
    search_terms, sections, specify_lines, use_total_lines, total_lines = data_json


    # Read the file
    file_content = file.read().decode('utf-8') #File content is in utf-8
    Lines = file_content.splitlines()
    

    # Create a new document
    document = Document()

    # Iterate over the search terms and find the corresponding sections
    for term in search_terms:
        lineNo = 0
        termLineNo = []
        termsNum = 0
        for line in Lines:
            if term in line:
                termLineNo.append(lineNo)
                termsNum += 1
                print(termsNum, term)
            lineNo += 1

    # Add the sections to the document
    for i in sections:
        section_lines = specify_lines[i-1].split()
        start_line = termLineNo[i-1]
        line_empty = 0

        if section_lines[0] == 'WHOLE' and use_total_lines == False:
            while line_empty == 0:
                if Lines[start_line] != "\n":
                    section = document.add_paragraph(Lines[start_line])
                    start_line += 1
                else:
                    line_empty = 1

        if section_lines[0] == 'WHOLE' and use_total_lines == True:
            for _ in range(total_lines - start_line + termLineNo[i-1]):
                section = document.add_paragraph(Lines[start_line])
                start_line += 1
                line_empty = 1
            else:
                start_line += 1
                line_empty = 1

        elif section_lines[0] == 'FIRST':
            line_count = -1
            while line_count < int(section_lines[1]):
                section = document.add_paragraph(Lines[start_line])
                start_line += 1
                line_count += 1

        elif section_lines[0] == 'LAST':
            line_count = -1
            document.add_paragraph(Lines[start_line])
            document.add_paragraph(Lines[start_line + 1])
            while line_count < int(section_lines[1]):
                section = document.add_paragraph(Lines[start_line+10])
                start_line += 1
                line_count += 1

        elif section_lines[0] == 'SPECIFIC':
            specific_lines = [int(l) for l in section_lines[1].split(",")]
            document.add_paragraph(Lines[start_line])
            for l in specific_lines:
                section = document.add_paragraph(Lines[start_line + l + 1])

    try:
        # Save the document
        final_location = Path("../../../esp/outputData/data_conversion.docx")
        final_location = final_location.resolve()
        document.save(final_location)
    except Exception as e:
        return f'Error saving document: {e}', 500

    return 'File processed successfully'
    


if __name__ == '__main__':
    app.run(debug=True)
from docx import Document

document = Document()

path = input("What is the file path? ")
readFile = open(path, "r")
Lines = readFile.readlines()

terms = []

def start():
    userInput = input("What section would you like to find? (enter exact name) ")
    searchTerm = userInput.upper()
    terms.append(searchTerm)

    proceed = 0

    while proceed == 0:
        moreYesNo = input("Is there another section you would like to find? (Y/N) ").upper()
        if moreYesNo == "Y":
            extraTerm = input("Please enter a specific section name: ")
            extraTermUpper = extraTerm.upper()
            terms.append(extraTermUpper)
        elif moreYesNo == "N":
            proceed = 1

def findTerms():
    for term in terms:
        lineNo = 0
        termLineNo = []
        lineEmpty = 0
        termsNum = 0
        for line in Lines:
            if term in line:
                termLineNo.append(lineNo)
                termsNum += 1
                print(termsNum, term)
            lineNo += 1
        sections = specifySection(term)
        print(sections)

        for i in sections:
            sectionLines = specifyLines(term, i)
            startLine = termLineNo[i-1]
            lineEmpty = 0
            use_total_lines = False

            if Lines[startLine+2] == "\n":
                use_total_lines=True
            else:
                use_total_lines = False

            if sectionLines[0] == 'WHOLE' and use_total_lines == False:
                while lineEmpty == 0:
                    if Lines[startLine] != "\n":
                        section = document.add_paragraph(Lines[startLine])
                        startLine += 1
                    else:
                        lineEmpty = 1

            if sectionLines[0] == 'WHOLE' and use_total_lines == True: #
                for _ in range(totalLine() - startLine + termLineNo[i-1]):
                    section = document.add_paragraph(Lines[startLine])
                    startLine += 1
                    lineEmpty = 1
                else:
                    startLine += 1
                    lineEmpty = 1

            if sectionLines[0] == 'FIRST':
                lineCount = 0
                while lineCount < int(sectionLines[1]):
                    section = document.add_paragraph(Lines[startLine])
                    startLine += 1
                    lineCount += 1

            elif sectionLines[0] == 'LAST':
                lineCount = 0
                title = document.add_paragraph(term)
                while lineCount < int(sectionLines[1]):
                    section = document.add_paragraph(Lines[startLine+10])
                    startLine += 1
                    lineCount += 1

            if sectionLines[0] == 'SPECIFIC':
                specific_lines = [int(l) for l in sectionLines[1].split(",")]
                document.add_paragraph(Lines[startLine])
                for l in specific_lines:
                    section = document.add_paragraph(Lines[startLine + l + 1])
                 
def closeAndSave():
    document.save("data_conversion.docx")
    readFile.close()

def specifySection(term):
    sectionInput = input("Which sections of " + term + " would you like to include? (e.g., 1,2) ")
    sections = [int(sectionNumber) for sectionNumber in sectionInput.split(',')]
    return sections

def specifyLines(term, sectionNo):
    linesInput = input("Would you like the WHOLE section, FIRST x lines, LAST x lines, or SPECIFIC x lines of " + term + " #" + str(sectionNo) + "? (e.g., FIRST 5) ")
    sectionLines = linesInput.split(" ")
    return sectionLines

def totalLine():
    total_lines = input("Please enter the total number of rows in the selected section: ")
    return int(total_lines)

def main():
    start()
    findTerms()
    closeAndSave()
    print(terms)


if __name__ == "__main__":
    main()
